"""Document generation tool -- create PDF or DOCX from text content.

Generates a well-formatted document from the user's request, uploads to S3,
and returns a DOC_URL:<url>:DOC_URL marker so the orchestrator delivers it as
a downloadable file via WhatsApp/Telegram.

Supports:
  - PDF via reportlab (already installed)
  - DOCX via python-docx (already installed)
"""
from __future__ import annotations

import io as _io
import logging
import os
import re
import uuid

import boto3

logger = logging.getLogger(__name__)

DOCUMENT_GEN_TOOL = {
    "toolSpec": {
        "name": "generate_document",
        "description": (
            "Create a formatted PDF or DOCX document from text content and send it to the user. "
            "Use this when the user asks you to: write a report, create a document, make a summary "
            "they can save, draft a letter/email/proposal, create a template, or anything they want "
            "as a downloadable file. "
            "Provide the full document content -- it will be formatted and sent as a file."
        ),
        "inputSchema": {
            "json": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Document title (shown as heading and filename)",
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Full document text. Use markdown-style formatting: "
                            "# Heading 1, ## Heading 2, **bold**, bullet lines starting with - or *. "
                            "Separate paragraphs with blank lines."
                        ),
                    },
                    "format": {
                        "type": "string",
                        "description": "Output format: \'pdf\' (default) or \'docx\'",
                        "default": "pdf",
                    },
                },
                "required": ["title", "content"],
            }
        },
    }
}


def _sanitize_filename(title: str) -> str:
    """Turn a title into a safe filename."""
    safe = re.sub(r"[^\w\s-]", "", title.lower())
    safe = re.sub(r"[\s-]+", "-", safe).strip("-")
    return safe[:60] or "document"


def _build_pdf(title: str, content: str) -> bytes:
    """Generate a clean PDF using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e"),
    )
    style_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"], fontSize=14, spaceAfter=4,
        textColor=colors.HexColor("#16213e"),
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=12, spaceAfter=3,
        textColor=colors.HexColor("#0f3460"),
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10.5, leading=15,
        spaceAfter=6, alignment=TA_LEFT,
    )
    style_bullet = ParagraphStyle(
        "Bullet", parent=style_body, leftIndent=14, bulletIndent=0,
    )

    story = []
    story.append(Paragraph(title, style_title))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 0.3 * cm))

    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            story.append(Spacer(1, 0.2 * cm))
            continue
        if line.startswith("# "):
            story.append(Paragraph(line[2:], style_h1))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], style_h2))
        elif line.startswith(("- ", "* ")):
            # Escape special chars for reportlab
            text = line[2:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Bold markers
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"&#8226; {text}", style_bullet))
        else:
            text = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(text, style_body))

    doc.build(story)
    return buf.getvalue()


def _build_docx(title: str, content: str) -> bytes:
    """Generate a DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_par = doc.add_heading(title, level=0)
    title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_bold_run(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_bold_run(p, line)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_bold_run(paragraph, text: str) -> None:
    """Add text to paragraph, handling **bold** markers."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


async def generate_document(title: str, content: str, format: str = "pdf") -> str:
    """Generate a PDF or DOCX, upload to S3, return DOC_URL marker."""
    bucket = os.environ.get("DATA_BUCKET", "")
    if not bucket:
        return "Document generation unavailable (DATA_BUCKET not configured)."

    fmt = (format or "pdf").lower().strip()
    if fmt not in ("pdf", "docx"):
        fmt = "pdf"

    try:
        if fmt == "pdf":
            file_bytes = _build_pdf(title, content)
            mime = "application/pdf"
            ext = "pdf"
        else:
            file_bytes = _build_docx(title, content)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"

        safe_name = _sanitize_filename(title)
        key = f"media/documents/{uuid.uuid4()}-{safe_name}.{ext}"

        s3_region = os.environ.get("AWS_REGION", "ap-southeast-1")
        s3 = boto3.client("s3", region_name=s3_region)
        s3.put_object(
            Bucket=bucket, Key=key, Body=file_bytes,
            ContentType=mime,
            ContentDisposition=f'attachment; filename="{safe_name}.{ext}"',
        )

        url = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": bucket, "Key": key},
            ExpiresIn=86400,  # 24h
        )

        logger.info("Generated document key=%s format=%s", key, fmt)
        return f"DOC_URL:{url}:DOC_URL"

    except Exception as e:
        logger.error("Document generation failed: %s", e)
        return f"Sorry, I couldn't generate the document: {str(e)[:120]}"
